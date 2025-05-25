# titanic_model_pipeline.py

from sklearn.metrics import precision_recall_curve, PrecisionRecallDisplay
from sklearn.metrics import roc_curve, RocCurveDisplay
from sklearn.metrics import confusion_matrix, ConfusionMatrixDisplay
import matplotlib.pyplot as plt
import os
import torch
import torch.nn as nn
import torch.optim as optim
from sklearn.model_selection import train_test_split
from sklearn.preprocessing import StandardScaler, LabelEncoder
from sklearn.metrics import accuracy_score, f1_score, recall_score, precision_score
import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Inches

# ============================
# 1) Carga y preprocesamiento de datos con split 70/15/15
# ============================

# 1.1 Cargar desde Excel
df = pd.read_excel("titanic3.xls")  # Ruta al archivo Excel

# Reemplazar coma decimal y convertir a numérico
df['fare'] = df['fare'].astype(str).str.replace(',', '.').astype(float)
df['age'] = df['age'].astype(str).str.replace(',', '.')
df['age'] = pd.to_numeric(df['age'], errors='coerce')

# Seleccionar columnas y eliminar nulos
df = df[['pclass', 'sex', 'age', 'sibsp', 'parch', 'fare', 'survived']].dropna()

# Codificar 'sex'
df['sex'] = LabelEncoder().fit_transform(df['sex'])

# Separar X e y
X = df.drop('survived', axis=1).values.astype(np.float32)
y = df['survived'].astype(int).values

# Escalado completo sobre todo X antes de dividir
scaler = StandardScaler()
X = scaler.fit_transform(X)

# Dividir en train (70%), temp (30%)
X_train, X_temp, y_train, y_temp = train_test_split(
    X, y, test_size=0.30, random_state=42, stratify=y)

# Dividir temp en validación (15%) y test (15%)
X_val, X_test, y_val, y_test = train_test_split(
    X_temp, y_temp, test_size=0.50, random_state=42, stratify=y_temp)

# Convertir a tensores de PyTorch
X_train = torch.tensor(X_train)
y_train = torch.tensor(y_train).float().unsqueeze(1)
X_val = torch.tensor(X_val)
y_val = torch.tensor(y_val).float().unsqueeze(1)
X_test = torch.tensor(X_test)
y_test = torch.tensor(y_test).float().unsqueeze(1)

# ============================
# 2) Definición dinámica del modelo y función de inicialización de pesos
# ============================


class MLP(nn.Module):
    def __init__(self, input_size, hidden_layers, activation_fn):
        super(MLP, self).__init__()
        layers = []
        last_size = input_size
        for h in hidden_layers:
            layers.append(nn.Linear(last_size, h))
            layers.append(activation_fn())
            last_size = h
        layers.append(nn.Linear(last_size, 1))
        layers.append(nn.Sigmoid())
        self.net = nn.Sequential(*layers)

    def forward(self, x):
        return self.net(x)


def initialize_weights(model, init_type="random"):
    """
    Inicializa los pesos según init_type.
    - "random": deja la inicialización por defecto de PyTorch.
    - "xavier": aplica Xavier uniform a todas las capas Lineales.
    """
    if init_type.lower() == "xavier":
        for m in model.modules():
            if isinstance(m, nn.Linear):
                nn.init.xavier_uniform_(m.weight)
                if m.bias is not None:
                    nn.init.constant_(m.bias, 0.0)
    # Si es "random", no hace nada (PyTorch ya inicializa aleatoriamente).

# ============================
# 3) Función de entrenamiento y evaluación
# ============================


def train_and_evaluate(config, data_tensors):
    """
    Entrena (o carga) el modelo según la configuración dada, calcula métricas,
    devuelve:
      - res: dict con métricas + hiperparámetros
      - model: el modelo entrenado (o cargado)
      - train_losses, val_losses: listas de las pérdidas por época
    """
    (X_tr, y_tr, X_v, y_v, X_te, y_te) = data_tensors
    label = config['label']
    model_path = f"models/{label}.pt"

    # Inicializar listas de pérdidas
    train_losses = []
    val_losses = []

    # Si existe el modelo guardado, lo cargamos; si no, entrenamos
    if os.path.exists(model_path):
        model = MLP(X_tr.shape[1], config['hidden_layers'],
                    config['activation_fn'])
        model.load_state_dict(torch.load(model_path))
        model.eval()
    else:
        model = MLP(X_tr.shape[1], config['hidden_layers'],
                    config['activation_fn'])
        # Inicializar pesos
        initialize_weights(model, config['weight_init'])

        # Seleccionar optimizador
        if config['optimizer_name'].lower() == "adam":
            optimizer = optim.Adam(model.parameters(), lr=config['lr'])
        elif config['optimizer_name'].lower() == "rmsprop":
            optimizer = optim.RMSprop(model.parameters(), lr=config['lr'])
        else:
            raise ValueError(
                f"Optimizador {config['optimizer_name']} no soportado.")

        criterion = nn.BCELoss()

        # Bucle de entrenamiento
        for epoch in range(config['epochs']):
            model.train()
            optimizer.zero_grad()

            # Forward pass
            outputs = model(X_tr)
            loss = criterion(outputs, y_tr)

            # Backward pass y optimización
            loss.backward()
            optimizer.step()

            # Guardar la pérdida de entrenamiento
            train_losses.append(loss.item())

            # Evaluación en validación
            model.eval()
            with torch.no_grad():
                val_outputs = model(X_v)
                val_loss = criterion(val_outputs, y_v)
                val_losses.append(val_loss.item())
            model.train()

        # Guardar el modelo entrenado
        os.makedirs("models", exist_ok=True)
        torch.save(model.state_dict(), model_path)
        model.eval()

    # Evaluar métricas en train / val / test
    with torch.no_grad():
        # Predicciones sobre train
        preds_tr = (model(X_tr) > 0.5).float()
        acc_tr = accuracy_score(y_tr.numpy(), preds_tr.numpy())
        f1_tr = f1_score(y_tr.numpy(), preds_tr.numpy())
        rec_tr = recall_score(y_tr.numpy(), preds_tr.numpy())
        prec_tr = precision_score(y_tr.numpy(), preds_tr.numpy())

        # Predicciones sobre validación
        preds_v = (model(X_v) > 0.5).float()
        acc_v = accuracy_score(y_v.numpy(), preds_v.numpy())
        f1_v = f1_score(y_v.numpy(), preds_v.numpy())
        rec_v = recall_score(y_v.numpy(), preds_v.numpy())
        prec_v = precision_score(y_v.numpy(), preds_v.numpy())

        # Predicciones sobre test
        preds_te = (model(X_te) > 0.5).float()
        acc_te = accuracy_score(y_te.numpy(), preds_te.numpy())
        f1_te = f1_score(y_te.numpy(), preds_te.numpy())
        rec_te = recall_score(y_te.numpy(), preds_te.numpy())
        prec_te = precision_score(y_te.numpy(), preds_te.numpy())

    res = {
        'label': label,
        'hyperparams': {
            'Optimizador': config['optimizer_name'],
            'Tasa de aprendizaje': config['lr'],
            'Pesos iniciales': config['weight_init'],
            'Épocas': config['epochs']
        },
        'metrics': {
            'Train': {'Accuracy': acc_tr, 'F1': f1_tr, 'Recall': rec_tr, 'Precision': prec_tr},
            'Val':   {'Accuracy': acc_v,  'F1': f1_v,  'Recall': rec_v,  'Precision': prec_v},
            'Test':  {'Accuracy': acc_te, 'F1': f1_te, 'Recall': rec_te, 'Precision': prec_te}
        }
    }

    return res, model, train_losses, val_losses

# ============================
# 4) Definir múltiples configuraciones paramétricas
# ============================


configs = [
    {
        'label': "modelo_1capa_Adam_lr1e-3_rand",
        'hidden_layers': [128] * 1,  # 1 capa de tamaño 128
        'activation_fn': nn.ReLU,
        'optimizer_name': "Adam",
        'lr': 1e-3,
        'weight_init': "random",
        'epochs': 100
    },
    {
        'label': "modelo_5capa_Adam_lr1e-3_rand",
        'hidden_layers': [128] * 5,  # 5 capas de tamaño 128
        'activation_fn': nn.ReLU,
        'optimizer_name': "Adam",
        'lr': 1e-3,
        'weight_init': "random",
        'epochs': 100
    },
    {
        'label': "modelo_8capa_Adam_lr1e-3_rand",
        'hidden_layers': [128] * 8,  # 8 capas de tamaño 128
        'activation_fn': nn.ReLU,
        'optimizer_name': "Adam",
        'lr': 1e-3,
        'weight_init': "random",
        'epochs': 100
    },
    {
        'label': "modelo_100capa_Adam_lr1e-3_rand",
        'hidden_layers': [128] * 100,  # 100 capas de tamaño 128
        'activation_fn': nn.ReLU,
        'optimizer_name': "Adam",
        'lr': 1e-3,
        'weight_init': "random",
        'epochs': 100
    }
]

# ============================
# Funciones auxiliares para gráficos y reportes
# ============================


def save_confusion_matrix(y_true, y_pred, labels, model_name, stage):
    cm = confusion_matrix(y_true, y_pred, normalize='true')
    disp = ConfusionMatrixDisplay(confusion_matrix=cm, display_labels=labels)
    disp.plot(cmap=plt.cm.Blues)
    plt.title(f'Matriz de Confusión Normalizada - {model_name} ({stage})')
    os.makedirs('reports', exist_ok=True)
    filepath = f'reports/confusion_matrix_{model_name}_{stage}.png'
    plt.savefig(filepath)
    plt.close()
    return filepath


def save_roc_curve(y_true, y_scores, model_name, stage):
    fpr, tpr, _ = roc_curve(y_true, y_scores)
    RocCurveDisplay(fpr=fpr, tpr=tpr).plot()
    plt.title(f'Curva ROC - {model_name} ({stage})')
    filepath = f'reports/roc_curve_{model_name}_{stage}.png'
    plt.savefig(filepath)
    plt.close()
    return filepath


def save_precision_recall_curve(y_true, y_scores, model_name, stage):
    precision, recall, _ = precision_recall_curve(y_true, y_scores)
    PrecisionRecallDisplay(precision=precision, recall=recall).plot()
    plt.title(f'Curva Precision-Recall - {model_name} ({stage})')
    filepath = f'reports/precision_recall_{model_name}_{stage}.png'
    plt.savefig(filepath)
    plt.close()
    return filepath


def save_learning_curve(train_losses, val_losses, model_name):
    plt.figure()
    plt.plot(train_losses, label='Entrenamiento')
    plt.plot(val_losses, label='Validación')
    plt.xlabel('Épocas')
    plt.ylabel('Pérdida')
    plt.title(f'Curva de Aprendizaje - {model_name}')
    plt.legend()
    filepath = f'reports/learning_curve_{model_name}.png'
    plt.savefig(filepath)
    plt.close()
    return filepath


def create_word_report(model_reports):
    document = Document()
    document.add_heading('Informe de Modelos - Titanic', 0)

    for report in model_reports:
        document.add_heading(f"Modelo: {report['model_name']}", level=1)

        # Incluimos Curva de Aprendizaje
        document.add_heading('Curva de Aprendizaje', level=2)
        document.add_picture(report['learning_curve'], width=Inches(5))

        for stage in ['validación', 'test']:
            document.add_heading(f"Etapa: {stage.capitalize()}", level=2)

            # Matriz de Confusión
            document.add_paragraph('Matriz de Confusión Normalizada:')
            document.add_picture(
                report['confusion_matrix'][stage], width=Inches(5))

            # Curva ROC
            document.add_paragraph('Curva ROC:')
            document.add_picture(report['roc_curve'][stage], width=Inches(5))

            # Curva Precision-Recall
            document.add_paragraph('Curva Precision-Recall:')
            document.add_picture(
                report['precision_recall'][stage], width=Inches(5))

    document.save('reports/Reporte_Modelos.docx')

# ============================
# 5) Loop principal: entrenar/cargar, recopilar métricas,
#    guardar reportes en Excel y generar reporte en Word
# ============================


all_results = []
model_reports = []
data_tensors = (X_train, y_train, X_val, y_val, X_test, y_test)

for idx, cfg in enumerate(configs, start=1):
    # (A) Entrenar o cargar + obtener métricas y pérdidas
    res, model_loaded, train_losses, val_losses = train_and_evaluate(
        cfg, data_tensors)

    hiper = res['hyperparams']
    met = res['metrics']

    # (B) Almacenar en all_results para Excel
    all_results.append({
        'Trat.': idx,
        'Optimizador': hiper['Optimizador'],
        'Tasa de aprendizaje': hiper['Tasa de aprendizaje'],
        'Pesos iniciales': hiper['Pesos iniciales'],
        'Épocas': hiper['Épocas'],
        'Exactitud entrenamiento': met['Train']['Accuracy'],
        'Exactitud validación': met['Val']['Accuracy'],
        'Exactitud prueba': met['Test']['Accuracy'],
        'Modelo': res['label'],
        'Test Accuracy': met['Test']['Accuracy'],
        'Test F1-Score': met['Test']['F1'],
        'Test Recall': met['Test']['Recall'],
        'Test Precision': met['Test']['Precision']
    })

    # (C) Generar gráficos para este modelo
    model_name = cfg['label']
    report = {
        'model_name': model_name,
        'confusion_matrix': {},
        'roc_curve': {},
        'precision_recall': {},
        'learning_curve': None
    }

    # Curva de Aprendizaje
    lc_path = save_learning_curve(train_losses, val_losses, model_name)
    report['learning_curve'] = lc_path

    # Para validación y test
    for stage, (X_stage, y_stage) in zip(['validación', 'test'], [(X_val, y_val), (X_test, y_test)]):
        with torch.no_grad():
            probs = model_loaded(X_stage).detach().numpy().flatten()
        y_pred = (probs > 0.5).astype(int)

        cm_path = save_confusion_matrix(
            y_stage.numpy().flatten(),
            y_pred,
            labels=['No', 'Sí'],
            model_name=model_name,
            stage=stage
        )
        report['confusion_matrix'][stage] = cm_path

        roc_path = save_roc_curve(
            y_stage.numpy().flatten(),
            probs,
            model_name=model_name,
            stage=stage
        )
        report['roc_curve'][stage] = roc_path

        pr_path = save_precision_recall_curve(
            y_stage.numpy().flatten(),
            probs,
            model_name=model_name,
            stage=stage
        )
        report['precision_recall'][stage] = pr_path

    model_reports.append(report)

# (D) Guardar Excel con métricas de todos los modelos
df_hyper = pd.DataFrame([
    {
        'Trat.': r['Trat.'],
        'Optimizador': r['Optimizador'],
        'Tasa de aprendizaje': r['Tasa de aprendizaje'],
        'Pesos iniciales': r['Pesos iniciales'],
        'Épocas': r['Épocas'],
        'Exactitud entrenamiento': r['Exactitud entrenamiento'],
        'Exactitud validación': r['Exactitud validación'],
        'Exactitud prueba': r['Exactitud prueba']
    }
    for r in all_results
])

df_modelos = pd.DataFrame([
    {
        'Modelo': r['Modelo'],
        'Accuracy': r['Test Accuracy'],
        'F1-Score': r['Test F1-Score'],
        'Recall': r['Test Recall'],
        'Precisión': r['Test Precision']
    }
    for r in all_results
])

os.makedirs("reports", exist_ok=True)
with pd.ExcelWriter("reports/reportes_modelos.xlsx") as writer:
    df_modelos.to_excel(writer, sheet_name="Metrics por Modelo", index=False)
    df_hyper.to_excel(
        writer, sheet_name="Hiperparam + Rendimiento", index=False)

print("✅ Se creó el reporte Excel en: reports/reportes_modelos.xlsx")

# (E) Generar el Word con todos los gráficos
create_word_report(model_reports)
print("✅ Se guardó el reporte Word en: reports/Reporte_Modelos.docx")
